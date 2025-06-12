# contract_generator/views.py

import pandas as pd
import docx
import re
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from .forms import DocumentTemplateForm, DataSourceForm
from .models import DocumentTemplate, DataSource

# --- Funciones Auxiliares ---
def extract_placeholders_from_docx(file_path):
    """Extrae campos como [campo] de un archivo .docx y devuelve una lista única."""
    try:
        doc = docx.Document(file_path)
        placeholders = set()
        # Expresión regular para encontrar texto entre corchetes
        pattern = re.compile(r'\[(.*?)\]')

        # Buscar en párrafos
        for para in doc.paragraphs:
            found = pattern.findall(para.text)
            for item in found:
                placeholders.add(item.strip())

        # Buscar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        found = pattern.findall(para.text)
                        for item in found:
                            placeholders.add(item.strip())
        return list(placeholders)
    except Exception as e:
        print(f"Error extrayendo placeholders: {e}")
        return []

def extract_headers_from_xlsx(file_path):
    """Extrae los encabezados de la primera fila de un archivo .xlsx."""
    try:
        # Usamos `openpyxl` como motor para leer .xlsx
        df = pd.read_excel(file_path, engine='openpyxl')
        return list(df.columns)
    except Exception as e:
        print(f"Error extrayendo headers: {e}")
        return []

# --- Vista Principal ---
@login_required # Este decorador asegura que solo usuarios logueados puedan acceder
def dashboard_view(request):
    # Procesar la subida de la plantilla de Word
    if 'submit_template' in request.POST:
        template_form = DocumentTemplateForm(request.POST, request.FILES)
        if template_form.is_valid():
            # Guardamos el formulario, pero sin confirmar en la BD todavía
            template = template_form.save(commit=False)
            template.owner = request.user # Asignamos el usuario actual como dueño
            template.save() # Ahora sí, guardamos en la BD

            # Extraemos los placeholders del archivo recién subido
            placeholders = extract_placeholders_from_docx(template.file.path)
            template.placeholders = placeholders # Guardamos la lista en el modelo
            template.save() # Volvemos a guardar para actualizar el campo

            return redirect('dashboard') # Redirigimos a la misma página
    else:
        template_form = DocumentTemplateForm()

    # Procesar la subida de la base de datos Excel
    if 'submit_datasource' in request.POST:
        datasource_form = DataSourceForm(request.POST, request.FILES)
        if datasource_form.is_valid():
            datasource = datasource_form.save(commit=False)
            datasource.owner = request.user
            datasource.save()

            # Extraemos los headers del archivo recién subido
            headers = extract_headers_from_xlsx(datasource.file.path)
            datasource.headers = headers # Guardamos la lista en el modelo
            datasource.save()

            return redirect('dashboard') # Redirigimos
    else:
        datasource_form = DataSourceForm()

    # Obtenemos los archivos ya subidos por el usuario para listarlos
    user_templates = DocumentTemplate.objects.filter(owner=request.user)
    user_datasources = DataSource.objects.filter(owner=request.user)

    context = {
        'template_form': template_form,
        'datasource_form': datasource_form,
        'user_templates': user_templates,
        'user_datasources': user_datasources,
    }
    return render(request, 'contract_generator/dashboard.html', context)

from .models import MappingProject
from .forms import MappingProjectForm
from django.shortcuts import get_object_or_404


@login_required
def create_mapping_project_view(request):
    if request.method == 'POST':
        # Le pasamos el usuario al formulario para que filtre los desplegables
        form = MappingProjectForm(request.POST, user=request.user)
        if form.is_valid():
            project = form.save(commit=False)
            project.owner = request.user
            project.save()
            # Una vez creado el proyecto, redirigimos a la página de mapeo de campos
            return redirect('map_fields', project_id=project.id)
    else:
        # Le pasamos el usuario para que el form se inicialice con los archivos correctos
        form = MappingProjectForm(user=request.user)

    # Obtenemos los proyectos existentes para listarlos
    existing_projects = MappingProject.objects.filter(owner=request.user)

    context = {
        'form': form,
        'projects': existing_projects
    }
    return render(request, 'contract_generator/create_mapping_project.html', context)


@login_required
def map_fields_view(request, project_id):
    # Obtenemos el proyecto o mostramos un error 404 si no existe o no es del usuario
    project = get_object_or_404(MappingProject, id=project_id, owner=request.user)

    if request.method == 'POST':
        # Borramos los mapeos antiguos para este proyecto por si el usuario está re-mapeando
        project.mappings.all().delete()

        # Los datos del formulario vienen del template, donde los campos del Word son las 'keys'
        for placeholder in project.template.placeholders:
            # El 'name' de cada <select> en el HTML será el nombre del placeholder
            selected_header = request.POST.get(placeholder)
            if selected_header: # Si el usuario seleccionó una columna para este campo
                # Creamos el objeto FieldMapping que guarda la relación
                project.mappings.create(
                    template_placeholder=placeholder,
                    data_header=selected_header
                )
        # Podríamos redirigir al dashboard o a una página de "generar documentos"
        return redirect('dashboard') 

    context = {
        'project': project,
        # Pasamos los campos detectados de la plantilla y las columnas de la base de datos
        'placeholders': project.template.placeholders,
        'headers': project.data_source.headers,
    }
    return render(request, 'contract_generator/map_fields.html', context)