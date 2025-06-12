# generator/views.py

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from .models import PlantillaContrato
from .models import FuenteDeDatos
import pandas as pd
import docx
import re
import os


def extraer_placeholders(ruta_archivo_docx):
    """
    Analiza un archivo .docx y extrae todas las etiquetas entre corchetes.
    """
    try:
        documento = docx.Document(ruta_archivo_docx)
        placeholders = set()
        regex = r"\[(.*?)\]"
        for p in documento.paragraphs:
            encontrados = re.findall(regex, p.text)
            for placeholder in encontrados:
                placeholders.add(placeholder.strip())
        for table in documento.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        encontrados = re.findall(regex, p.text)
                        for placeholder in encontrados:
                            placeholders.add(placeholder.strip())
        return list(placeholders)
    except Exception as e:
        print(f"Error crítico al analizar el archivo DOCX: {e}")
        return None


@login_required
def gestion_plantillas(request):
    context = {
        'plantillas': PlantillaContrato.objects.filter(propietario=request.user),
        'error': None
    }

    if request.method == 'POST':
        nombre_plantilla = request.POST.get('nombre_plantilla')
        archivo_plantilla = request.FILES.get('archivo_plantilla')

        # 1. Validar que ambos campos están presentes
        if not nombre_plantilla or not archivo_plantilla:
            context['error'] = "Error: Debes proporcionar un nombre y seleccionar un archivo."
            return render(request, 'generator/gestion_plantillas.html', context)

        # 2. Validar la extensión del archivo de forma segura
        nombre_archivo = archivo_plantilla.name
        _ , extension = os.path.splitext(nombre_archivo)

        if extension.lower() != '.docx':
            context['error'] = (
                "Formato de archivo no válido. "
                "Solo se aceptan archivos Word modernos (.docx). "
                "Si tu archivo es .doc, por favor ábrelo en Word y guárdalo como .docx."
            )
            return render(request, 'generator/gestion_plantillas.html', context)

        nueva_plantilla = PlantillaContrato(
            propietario=request.user,
            nombre=nombre_plantilla,
            archivo=archivo_plantilla
        )
        nueva_plantilla.save()

        placeholders_extraidos = extraer_placeholders(nueva_plantilla.archivo.path)
        
        if placeholders_extraidos is None:
            context['error'] = "Error: El archivo .docx parece estar corrupto o no se puede leer."
            # Borramos el registro y el archivo si no se puede procesar
            nueva_plantilla.delete() 
            return render(request, 'generator/gestion_plantillas.html', context)

        nueva_plantilla.placeholders = placeholders_extraidos
        nueva_plantilla.save()
        
        return redirect('generator:gestion_plantillas')

    return render(request, 'generator/gestion_plantillas.html', context)


def extraer_columnas_excel(ruta_archivo_excel):
    """
    Analiza un archivo Excel y extrae los nombres de las columnas de la primera fila.
    """
    try:
        # Leemos solo la primera fila (nrows=0) para obtener los encabezados.
        # Esto es muy eficiente, no carga el archivo entero en memoria.
        df = pd.read_excel(ruta_archivo_excel, nrows=0)
        return df.columns.tolist()
    except Exception as e:
        # Si pandas no puede leerlo (corrupto, formato incorrecto, etc.)
        print(f"Error crítico al analizar el archivo Excel: {e}")
        return None # Devolvemos None para indicar un fallo

@login_required
def gestion_fuentes_datos(request):
    context = {
        'fuentes_datos': FuenteDeDatos.objects.filter(propietario=request.user),
        'error': None
    }

    if request.method == 'POST':
        nombre_fuente = request.POST.get('nombre_fuente')
        archivo_fuente = request.FILES.get('archivo_fuente')

        # 1. Validar que los campos están presentes
        if not nombre_fuente or not archivo_fuente:
            context['error'] = "Error: Debes proporcionar un nombre y seleccionar un archivo."
            return render(request, 'generator/gestion_fuentes_datos.html', context)

        # 2. Validar la extensión del archivo
        _ , extension = os.path.splitext(archivo_fuente.name)
        if extension.lower() != '.xlsx':
            context['error'] = "Formato de archivo no válido. Solo se aceptan archivos Excel (.xlsx)."
            return render(request, 'generator/gestion_fuentes_datos.html', context)

        # Si la validación es correcta, procedemos
        nueva_fuente = FuenteDeDatos(
            propietario=request.user,
            nombre=nombre_fuente,
            archivo=archivo_fuente
        )
        nueva_fuente.save()

        # Intentamos extraer las columnas
        columnas_extraidas = extraer_columnas_excel(nueva_fuente.archivo.path)
        
        if columnas_extraidas is None:
            context['error'] = "Error: El archivo .xlsx parece estar corrupto o no es un formato válido."
            nueva_fuente.delete()
            return render(request, 'generator/gestion_fuentes_datos.html', context)

        nueva_fuente.columnas = columnas_extraidas
        nueva_fuente.save()
        
        return redirect('generator:gestion_fuentes_datos')

    return render(request, 'generator/gestion_fuentes_datos.html', context)

